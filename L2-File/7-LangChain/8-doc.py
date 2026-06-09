# 8-doc.py - 完整修复版
from langchain_text_splitters import RecursiveCharacterTextSplitter 
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings 
from langchain_core.runnables import RunnableLambda, RunnablePassthrough
from langchain_core.prompts import ChatPromptTemplate
from langchain_ollama import ChatOllama
from langchain_core.output_parsers import StrOutputParser
from langchain_core.documents import Document as LangChainDocument
import zipfile
from pathlib import Path
from langchain_community.document_loaders import Docx2txtLoader

def load_docx_smart(file_name: str) -> list:
    """智能加载 DOCX 文件"""
    # 🔴 修复：正确的路径计算
    current_dir = Path(__file__).resolve().parent  # e:\py-file\L2-File\7-LangChain
    project_dir = current_dir.parent  # e:\py-file\L2-File
    data_dir = project_dir / "Data"  # e:\py-file\L2-File\Data
    file_path = data_dir / file_name  # e:\py-file\L2-File\Data\datasecurity.docx
    
    # 🔴 添加路径调试信息
    print(f"📁 当前脚本：{current_dir}")
    print(f"📁 项目目录：{project_dir}")
    print(f"📁 数据目录：{data_dir}")
    print(f"📄 文件路径：{file_path}")
    
    # 🔴 检查文件是否存在
    if not file_path.exists():
        print(f"⚠ 文件不存在：{file_path}")
        # 尝试备用路径
        alt_paths = [
            Path(__file__).resolve().parent.parent.parent / "Data" / file_name,
            Path(__file__).resolve().parent / "Data" / file_name,
            Path.cwd() / "Data" / file_name,
        ]
        for alt_path in alt_paths:
            if alt_path.exists():
                print(f"✓ 找到备用路径：{alt_path}")
                file_path = alt_path
                break
        else:
            raise FileNotFoundError(f"无法找到文件：{file_name}")
    
    print(f"✓ 最终使用路径：{file_path}")

    # 检查是否为有效 ZIP 文件
    try:
        with zipfile.ZipFile(file_path, 'r') as zipf:
            zipf.testzip()
        print("✓ 文件是有效的 .docx 格式")
        
        loader = Docx2txtLoader(file_path=file_path)
        return loader.load()
        
    except zipfile.BadZipFile:
        print("⚠ 文件不是有效的 .docx 格式，使用备用加载器...")
        from docx import Document
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return [LangChainDocument(page_content=content, metadata={"source": file_path})]

llm = ChatOllama(model="qwen2.5:7b", temperature=0.5)
strparser = StrOutputParser()

embeddings = OllamaEmbeddings( 
    model="bge-m3:latest",
    base_url="http://127.0.0.1:11434"
)

# 🔴 调用函数
documents = load_docx_smart("datasecurity.docx")
print(f"文档数量：{len(documents)}")

text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
split_documents = text_splitter.split_documents(documents)

db = Chroma.from_documents(documents=split_documents, embedding=embeddings)

docs_find = RunnableLambda(db.similarity_search).bind(k=1)

message = """
仅使用提供的上下文回答下面的问题：
{question}
上下文：
{context}
"""
prompt_template = ChatPromptTemplate.from_messages(
    [
        ('human', message)  # 这个需要是human,而不是ai,可以是system,但效果没human好
    ]
)


chain = {"question": RunnablePassthrough(), "context": docs_find} | prompt_template | llm | strparser

# resp = chain.invoke("信息泄露有哪些风险")
user_question = str(input("请输入问题:"))
resp = chain.invoke(user_question)

print("resp:\n",resp)